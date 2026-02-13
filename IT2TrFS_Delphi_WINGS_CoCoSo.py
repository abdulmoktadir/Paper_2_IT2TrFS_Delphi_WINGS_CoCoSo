import streamlit as st
import numpy as np
import pandas as pd
import graphviz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import matplotlib.pyplot as plt
import re

# ============================================================
# IT2TrFS core (your WINGS code kept as-is, only reorganized)
# ============================================================

# Define linguistic terms for IT2TrFS
LINGUISTIC_TERMS = {
    "strength": {
        "VLR": ((0, 0.1, 0.1, 0.1, 1, 1), (0.0, 0.1, 0.1, 0.05, 0.9, 0.9)),
        "LR":  ((0.2, 0.3, 0.3, 0.4, 1, 1), (0.25, 0.3, 0.3, 0.35, 0.9, 0.9)),
        "MR":  ((0.4, 0.5, 0.5, 0.6, 1, 1), (0.45, 0.5, 0.5, 0.55, 0.9, 0.9)),
        "HR":  ((0.6, 0.7, 0.7, 0.8, 1, 1), (0.65, 0.7, 0.7, 0.75, 0.9, 0.9)),
        "VHR": ((0.8, 0.9, 0.9, 1,   1, 1), (0.85, 0.90,0.90,0.95, 0.9, 0.9)),
    },
    "influence": {
        "ELI": ((0,   0.1, 0.1, 0.2, 1, 1), (0.05,0.1, 0.1, 0.15,0.9,0.9)),
        "VLI": ((0.1, 0.2, 0.2, 0.35,1, 1), (0.15,0.2, 0.2, 0.3, 0.9,0.9)),
        "LI":  ((0.2, 0.35,0.35,0.5, 1, 1), (0.25,0.35,0.35,0.45,0.9,0.9)),
        "MI":  ((0.35,0.5, 0.5, 0.65,1, 1), (0.40,0.5, 0.5, 0.6, 0.9,0.9)),
        "HI":  ((0.5, 0.65,0.65,0.8, 1, 1), (0.55,0.65,0.65,0.75,0.9,0.9)),
        "VHI": ((0.65,0.80,0.80,0.9, 1, 1), (0.7, 0.8, 0.8, 0.85,0.9,0.9)),
        "EHI": ((0.8, 0.9, 0.9, 1,   1, 1), (0.85,0.9, 0.9, 0.95,0.9,0.9)),
    }
}

# Full forms
FULL_FORMS = {
    "VLR": "Very Low Relevance",
    "LR":  "Low Relevance",
    "MR":  "Medium Relevance",
    "HR":  "High Relevance",
    "VHR": "Very High Relevance",
    "ELI": "Extremely Low Influence",
    "VLI": "Very Low Influence",
    "LI":  "Low Influence",
    "MI":  "Medium Influence",
    "HI":  "High Influence",
    "VHI": "Very High Influence",
    "EHI": "Extremely High Influence",
}

def format_it2(it2):
    u, l = it2
    return f"(({u[0]:.6f},{u[1]:.6f},{u[2]:.6f},{u[3]:.6f};{u[4]:.1f},{u[5]:.1f}), ({l[0]:.6f},{l[1]:.6f},{l[2]:.6f},{l[3]:.6f};{l[4]:.1f},{l[5]:.1f}))"

def add_it2(A, B):
    Au, Al = A; Bu, Bl = B
    new_u = (Au[0]+Bu[0], Au[1]+Bu[1], Au[2]+Bu[2], Au[3]+Bu[3], min(Au[4],Bu[4]), min(Au[5],Bu[5]))
    new_l = (Al[0]+Bl[0], Al[1]+Bl[1], Al[2]+Bl[2], Al[3]+Bl[3], min(Al[4],Bl[4]), min(Al[5],Bl[5]))
    return (new_u, new_l)

def sub_it2(A, B):
    Au, Al = A; Bu, Bl = B
    new_u = (Au[0]-Bu[0], Au[1]-Bu[1], Au[2]-Bu[2], Au[3]-Bu[3], min(Au[4],Bu[4]), min(Au[5],Bu[5]))
    new_l = (Al[0]-Bl[0], Al[1]-Bl[1], Al[2]-Bl[2], Al[3]-Bl[3], min(Al[4],Bl[4]), min(Al[5],Bl[5]))
    return (new_u, new_l)

def mul_it2(A, B):
    Au, Al = A; Bu, Bl = B
    new_u = (Au[0]*Bu[0], Au[1]*Bu[1], Au[2]*Bu[2], Au[3]*Bu[3], min(Au[4],Bu[4]), min(Au[5],Bu[5]))
    new_l = (Al[0]*Bl[0], Al[1]*Bl[1], Al[2]*Bl[2], Al[3]*Bl[3], min(Al[4],Bl[4]), min(Al[5],Bl[5]))
    return (new_u, new_l)

def scalar_mul_it2(k, A):
    Au, Al = A
    new_u = (k*Au[0], k*Au[1], k*Au[2], k*Au[3], Au[4], Au[5])
    new_l = (k*Al[0], k*Al[1], k*Al[2], k*Al[3], Al[4], Al[5])
    return (new_u, new_l)

def zero_it2():
    return ((0,0,0,0,1,1), (0,0,0,0,0.9,0.9))

def defuzz_it2(A):
    Au, Al = A
    return (Au[0]+Au[1]+Au[2]+Au[3]+Al[0]+Al[1]+Al[2]+Al[3]) / 8

def identity_it2(n):
    I_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        I_mat[i][i] = ((1,1,1,1,1,1), (1,1,1,1,1,1))
    return I_mat

def compute_total_relation_matrix(normalized_matrix):
    n = len(normalized_matrix)
    I = identity_it2(n)

    # Convert normalized_matrix to 4D array for parameter-wise computation
    Z_4d = np.zeros((2, 2, n, n, 4))
    for i in range(n):
        for j in range(n):
            Au, Al = normalized_matrix[i][j]
            Z_4d[0, 0, i, j, :] = Au[:4]          # UMF a,b,c,d
            Z_4d[0, 1, i, j, :2] = Au[4:]         # UMF heights
            Z_4d[1, 0, i, j, :] = Al[:4]          # LMF a,b,c,d
            Z_4d[1, 1, i, j, :2] = Al[4:]         # LMF heights

    # Compute T for each parameter (a,b,c,d only)
    for i in range(2):  # UMF/LMF
        for k in range(4):
            Z_component = Z_4d[i, 0, :, :, k]
            try:
                T_component = Z_component @ np.linalg.pinv(np.eye(n) - Z_component)
            except np.linalg.LinAlgError:
                T_component = np.zeros((n, n))
            Z_4d[i, 0, :, :, k] = T_component

    # Reconstruct IT2TrFS matrix
    T = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            T[i][j] = (
                (Z_4d[0,0,i,j,0], Z_4d[0,0,i,j,1], Z_4d[0,0,i,j,2], Z_4d[0,0,i,j,3],
                 Z_4d[0,1,i,j,0], Z_4d[0,1,i,j,1]),
                (Z_4d[1,0,i,j,0], Z_4d[1,0,i,j,1], Z_4d[1,0,i,j,2], Z_4d[1,0,i,j,3],
                 Z_4d[1,1,i,j,0], Z_4d[1,1,i,j,1]),
            )
    return T

def calculate_TI_TR(T):
    n = len(T)
    TI = [zero_it2() for _ in range(n)]
    TR = [zero_it2() for _ in range(n)]
    for i in range(n):
        for j in range(n):
            TI[i] = add_it2(TI[i], T[i][j])
            TR[j] = add_it2(TR[j], T[i][j])
    return TI, TR

def wings_method_experts(strengths_list, influence_matrices_list, weights=None):
    n = len(strengths_list[0])
    num_experts = len(strengths_list)
    if weights is None:
        weights = [1.0/num_experts]*num_experts

    # weighted average SIDRM
    avg_sidrm = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for exp in range(num_experts):
        w = weights[exp]
        for i in range(n):
            avg_sidrm[i][i] = add_it2(avg_sidrm[i][i], scalar_mul_it2(w, strengths_list[exp][i]))
            for j in range(n):
                if i != j:
                    avg_sidrm[i][j] = add_it2(avg_sidrm[i][j], scalar_mul_it2(w, influence_matrices_list[exp][i][j]))

    # normalization scalar s (sum of all a,b,c,d across UMF+LMF)
    s = 0.0
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            s += (Au[0]+Au[1]+Au[2]+Au[3]+Al[0]+Al[1]+Al[2]+Al[3])

    Z_mat = [[zero_it2() for _ in range(n)] for _ in range(n)]
    for i in range(n):
        for j in range(n):
            Au, Al = avg_sidrm[i][j]
            new_u = (Au[0]/s if s else 0, Au[1]/s if s else 0, Au[2]/s if s else 0, Au[3]/s if s else 0, Au[4], Au[5])
            new_l = (Al[0]/s if s else 0, Al[1]/s if s else 0, Al[2]/s if s else 0, Al[3]/s if s else 0, Al[4], Al[5])
            Z_mat[i][j] = (new_u, new_l)

    T_mat = compute_total_relation_matrix(Z_mat)
    TI, TR = calculate_TI_TR(T_mat)

    engagement = [add_it2(TI[i], TR[i]) for i in range(n)]
    role       = [sub_it2(TI[i], TR[i]) for i in range(n)]

    TI_defuzz = np.array([defuzz_it2(TI[i]) for i in range(n)])
    TR_defuzz = np.array([defuzz_it2(TR[i]) for i in range(n)])
    engagement_defuzz = np.array([defuzz_it2(engagement[i]) for i in range(n)])
    role_defuzz = np.array([defuzz_it2(role[i]) for i in range(n)])

    return {
        'average_sidrm': avg_sidrm,
        'scaling_factor': s,
        'normalized_matrix': Z_mat,
        'total_matrix': T_mat,
        'total_impact': TI,
        'total_receptivity': TR,
        'engagement': engagement,
        'role': role,
        'total_impact_defuzz': TI_defuzz,
        'total_receptivity_defuzz': TR_defuzz,
        'engagement_defuzz': engagement_defuzz,
        'role_defuzz': role_defuzz
    }

def format_it2_df(mat, index, columns):
    df = pd.DataFrame(index=index, columns=columns)
    for i in range(len(index)):
        for j in range(len(columns)):
            df.iloc[i, j] = format_it2(mat[i][j])
    return df

def generate_flowchart_for_expert(expert_data, component_names, expert_idx=None):
    graph = graphviz.Digraph(comment=f'IT2TrFS WINGS Flowchart - Expert {expert_idx+1}' if expert_idx is not None else 'IT2TrFS WINGS Flowchart')
    graph.attr(rankdir='TD', size='8,8')

    for comp_idx, comp_name in enumerate(component_names):
        strength = expert_data['strengths_linguistic'][comp_idx]
        label = f"{comp_name} ({strength})"
        graph.node(comp_name, label=label, shape='box', style='rounded,filled', fillcolor='lightblue', fontsize='12')

    for from_idx, from_comp in enumerate(component_names):
        for to_idx, to_comp in enumerate(component_names):
            if from_idx == to_idx:
                continue
            influence = expert_data['influence_matrix_linguistic'][from_idx][to_idx]
            if influence != "ELI":
                graph.edge(from_comp, to_comp, label=influence)

    return graph

def add_dataframe_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns)+1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    for i, col in enumerate(df.columns):
        hdr_cells[i+1].text = str(col)
    for i, index in enumerate(df.index):
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        for j, col in enumerate(df.columns):
            row_cells[j+1].text = str(df.iloc[i, j])
    doc.add_paragraph()

def create_word_report(results, component_names, n_experts=1, expert_weights=None):
    doc = Document()
    title = doc.add_heading('IT2TrFS WINGS Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Number of experts: {n_experts}")

    if expert_weights and n_experts > 1:
        doc.add_paragraph("Expert weights: " + ", ".join([f"Expert {i+1}: {w:.2f}" for i, w in enumerate(expert_weights)]))

    comp_para = doc.add_paragraph("Components analyzed: ")
    for i, name in enumerate(component_names):
        comp_para.add_run(f"{i+1}. {name}  ")

    doc.add_heading('Impact, Receptivity, Engagement, and Role Results', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Total Impact (TI)'
    hdr[2].text = 'Total Receptivity (TR)'
    hdr[3].text = 'Engagement (TI+TR)'
    hdr[4].text = 'Role (TI-TR)'

    for i, name in enumerate(component_names):
        r = table.add_row().cells
        r[0].text = name
        r[1].text = f"{results['total_impact_defuzz'][i]:.6f}"
        r[2].text = f"{results['total_receptivity_defuzz'][i]:.6f}"
        r[3].text = f"{results['engagement_defuzz'][i]:.6f}"
        r[4].text = f"{results['role_defuzz'][i]:.6f}"

    doc.add_heading('Component Classification', level=1)
    class_table = doc.add_table(rows=1, cols=3)
    class_table.style = 'Table Grid'
    hdr = class_table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Type'
    hdr[2].text = 'Role (TI-TR)'

    for i, name in enumerate(component_names):
        status = "Cause" if results['role_defuzz'][i] > 0 else "Effect"
        r = class_table.add_row().cells
        r[0].text = name
        r[1].text = status
        r[2].text = f"{results['role_defuzz'][i]:.6f}"

    doc.add_heading('Matrices', level=1)
    doc.add_heading('Average SIDRM', level=2)
    add_dataframe_to_doc(doc, format_it2_df(results['average_sidrm'], component_names, component_names))
    doc.add_heading('Normalized Matrix Z', level=2)
    add_dataframe_to_doc(doc, format_it2_df(results['normalized_matrix'], component_names, component_names))
    doc.add_heading('Total Matrix T (IT2TrFS)', level=2)
    add_dataframe_to_doc(doc, format_it2_df(results['total_matrix'], component_names, component_names))

    doc.add_heading('Interpretation of Results', level=1)
    doc.add_paragraph("Total Impact (TI) represents the outgoing influence of a component.")
    doc.add_paragraph("Total Receptivity (TR) represents the incoming influence on a component.")
    doc.add_paragraph("Engagement (TI+TR) indicates the overall involvement of a component in the system.")
    doc.add_paragraph("Role (TI-TR) indicates cause/effect: positive = Cause, negative = Effect.")
    return doc

def get_word_download_link(doc):
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    b64 = base64.b64encode(file_stream.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="it2trfs_wings_analysis_report.docx">Download Word Report</a>'


# ============================================================
# IT2TrFS-CoCoSo (Excel-driven integration)
# Reads your sample Excel and presents:
# - linguistic scale (sheet1)
# - expert linguistic decision matrices (EX-1..EX-n)
# - weights (sheet2)
# - normalized matrices per alternative (sheet2)
# - Sbi/Pbi and final K + Rank (sheet2)
# ============================================================

def _safe_str(x):
    return "" if pd.isna(x) else str(x)

def _find_cell(df, target):
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            if df.iat[i, j] == target:
                return (i, j)
    return None

def _parse_it2_from_row(row_dict):
    # expects keys: a,b,c,d,uh1,uh2,e,f,g,h,lh1,lh2
    u = (float(row_dict["a"]), float(row_dict["b"]), float(row_dict["c"]), float(row_dict["d"]),
         float(row_dict["uh1"]), float(row_dict["uh2"]))
    l = (float(row_dict["e"]), float(row_dict["f"]), float(row_dict["g"]), float(row_dict["h"]),
         float(row_dict["lh1"]), float(row_dict["lh2"]))
    return (u, l)

def parse_cocoso_excel(xlsx_bytes_or_path):
    xls = pd.ExcelFile(xlsx_bytes_or_path)

    # --- sheet 1: input + linguistic table ---
    sh1 = pd.read_excel(xlsx_bytes_or_path, sheet_name=xls.sheet_names[0], header=None)

    # linguistic scale table
    loc = _find_cell(sh1, "Linguistic Attribute")
    if not loc:
        raise ValueError("Could not find 'Linguistic Attribute' table in sheet 1.")
    r0, c0 = loc

    # expected columns: [Linguistic Attribute, IT-2FNumber, (string), a,b,c,d,uh1,uh2,e,f,g,h,lh1,lh2]
    scale_rows = []
    for r in range(r0 + 1, r0 + 20):
        term = _safe_str(sh1.iat[r, c0]).strip()
        code = _safe_str(sh1.iat[r, c0 + 1]).strip()
        if not term or not code:
            continue
        a = sh1.iat[r, c0 + 3]; b = sh1.iat[r, c0 + 4]; c = sh1.iat[r, c0 + 5]; d = sh1.iat[r, c0 + 6]
        uh1 = sh1.iat[r, c0 + 7]; uh2 = sh1.iat[r, c0 + 8]
        e = sh1.iat[r, c0 + 9]; f = sh1.iat[r, c0 + 10]; g = sh1.iat[r, c0 + 11]; h = sh1.iat[r, c0 + 12]
        lh1 = sh1.iat[r, c0 + 13]; lh2 = sh1.iat[r, c0 + 14]
        if any(pd.isna(v) for v in [a,b,c,d,uh1,uh2,e,f,g,h,lh1,lh2]):
            continue
        scale_rows.append({
            "Linguistic Attribute": term,
            "Code": code,
            "a": float(a), "b": float(b), "c": float(c), "d": float(d),
            "uh1": float(uh1), "uh2": float(uh2),
            "e": float(e), "f": float(f), "g": float(g), "h": float(h),
            "lh1": float(lh1), "lh2": float(lh2),
            "IT2TrFS": _parse_it2_from_row({
                "a": a, "b": b, "c": c, "d": d, "uh1": uh1, "uh2": uh2,
                "e": e, "f": f, "g": g, "h": h, "lh1": lh1, "lh2": lh2
            })
        })

    scale_df = pd.DataFrame(scale_rows)
    code_to_it2 = {row["Code"]: row["IT2TrFS"] for _, row in scale_df.iterrows()}

    # expert blocks: detect EX-* in row 0
    ex_cols = []
    for j in range(sh1.shape[1]):
        v = sh1.iat[0, j]
        if isinstance(v, str) and v.strip().startswith("EX-"):
            ex_cols.append((v.strip(), j))

    # criteria in col 1, optimization in col 2 starting row 4 until blank
    criteria = []
    optm = []
    r = 4
    while r < sh1.shape[0]:
        c_name = _safe_str(sh1.iat[r, 1]).strip()
        if not c_name:
            break
        criteria.append(c_name)
        optm.append(_safe_str(sh1.iat[r, 2]).strip())
        r += 1

    # each EX block has 7 alternatives columns (VT1..VT7) in this template
    experts = []
    for ex_name, start_col in ex_cols:
        # alt codes on row 3, names on row 2
        alt_codes = []
        alt_names = []
        for j in range(start_col, start_col + 7):
            code = _safe_str(sh1.iat[3, j]).strip()
            name = _safe_str(sh1.iat[2, j]).strip()
            if code:
                alt_codes.append(code.replace(" ", ""))
                alt_names.append(name if name else code)
        # decision matrix linguistic codes
        mat = pd.DataFrame(index=criteria, columns=alt_codes, dtype=str)
        for i_c, c_name in enumerate(criteria):
            rr = 4 + i_c
            for k, ac in enumerate(alt_codes):
                cc = start_col + k
                mat.loc[c_name, ac] = _safe_str(sh1.iat[rr, cc]).strip()
        experts.append({
            "expert_label": ex_name,
            "alt_codes": alt_codes,
            "alt_names": alt_names,
            "criteria": criteria,
            "optimization": optm,
            "linguistic_matrix": mat
        })

    # --- sheet 2: computed pipeline + weights + final ranking ---
    sh2 = pd.read_excel(xlsx_bytes_or_path, sheet_name=xls.sheet_names[1], header=None)

    # weights appear aligned with rows where col1 is CSF1.. (in your sample rows 20..31) and col86 has weights
    # Find a "Weight" header:
    w_loc = _find_cell(sh2, "Weight")
    if not w_loc:
        raise ValueError("Could not find 'Weight' column in sheet 2.")
    _, w_col = w_loc

    # Find block where first CSF appears near the weights; in sample it's rows 20..31
    csf_rows = []
    for rr in range(sh2.shape[0]):
        v = _safe_str(sh2.iat[rr, 1]).strip()
        if re.fullmatch(r"CSF\d+", v):
            csf_rows.append(rr)

    # take the contiguous run starting at the first CSF after the "CSFs a b c ..." header near row 19
    # simplest: take the 12 CSFs with weights present (non-NaN)
    weight_rows = [rr for rr in csf_rows if not pd.isna(sh2.iat[rr, w_col])]
    # pick first 12
    weight_rows = sorted(weight_rows)[:12]

    crit2 = [_safe_str(sh2.iat[rr, 1]).strip() for rr in weight_rows]
    opt2  = [_safe_str(sh2.iat[rr, 0]).strip() for rr in weight_rows]
    weights = [float(sh2.iat[rr, w_col]) for rr in weight_rows]

    crit_weight_df = pd.DataFrame({"Criterion": crit2, "Optimization": opt2, "Weight": weights})

    # alternatives block headers: row where col0 is 'Optm.' and then several 'T1','T2',...
    # in sample it's row 18, but detect by scanning for 'Optm.' in col0
    header_r = None
    for rr in range(sh2.shape[0]):
        if _safe_str(sh2.iat[rr, 0]).strip() == "Optm.":
            header_r = rr
            break
    if header_r is None:
        raise ValueError("Could not find 'Optm.' header row in sheet 2.")

    alt_starts = []
    for j in range(sh2.shape[1]):
        v = _safe_str(sh2.iat[header_r, j]).strip()
        if re.fullmatch(r"T\d+", v):
            alt_starts.append((v, j))
    alt_codes2 = [a for a, _ in alt_starts]

    # normalized decision matrices (rows weight_rows, each alternative has 12 columns: a..lh2)
    # column layout in your sheet: at start: a b c d uh1 uh2 e f g h lh1 lh2 (12 cols)
    keys12 = ["a","b","c","d","uh1","uh2","e","f","g","h","lh1","lh2"]
    norm_mats = {}  # alt_code -> DataFrame(criteria x 1 cell IT2TrFS)
    for alt, start_c in alt_starts:
        # the normalized values are located at the same columns start_c..start_c+11 and rows weight_rows
        tmp = []
        for rr, c_name in zip(weight_rows, crit2):
            row = {k: sh2.iat[rr, start_c + idx] for idx, k in enumerate(keys12)}
            if any(pd.isna(row[k]) for k in keys12):
                it2 = None
            else:
                it2 = _parse_it2_from_row({k: float(row[k]) for k in keys12})
            tmp.append((c_name, it2))
        norm_mats[alt] = pd.DataFrame({"Criterion": [t[0] for t in tmp], "IT2TrFS": [t[1] for t in tmp]}).set_index("Criterion")

    # Sbi/Pbi and final K + Rank block (in your sample starts at row with 'Sbi' in col2)
    sbi_loc = _find_cell(sh2, "Sbi")
    if not sbi_loc:
        raise ValueError("Could not find 'Sbi' block in sheet 2.")
    sbi_r, _ = sbi_loc

    # rows with alternatives are sbi_r+2 .. (until blank)
    rows_alt = []
    for rr in range(sbi_r + 2, sh2.shape[0]):
        a = _safe_str(sh2.iat[rr, 1]).strip()
        if re.fullmatch(r"T\d+", a):
            rows_alt.append(rr)
        elif rows_alt:
            break

    # Crisp columns are in your sample: 27 (Crisp SBi) and 28 (Crisp PBi)
    # Kia/Kib/Kic/K/Rank are in 31..35
    final_rows = []
    for rr in rows_alt:
        alt = _safe_str(sh2.iat[rr, 1]).strip()
        # Technology names appear in col37 in sample
        tech = _safe_str(sh2.iat[rr, 37]).strip()
        final_rows.append({
            "Alt": alt,
            "Technology": tech if tech else alt,
            "Crisp_SBi": float(sh2.iat[rr, 27]),
            "Crisp_PBi": float(sh2.iat[rr, 28]),
            "Kia": float(sh2.iat[rr, 31]),
            "Kib": float(sh2.iat[rr, 32]),
            "Kic": float(sh2.iat[rr, 33]),
            "K": float(sh2.iat[rr, 34]),
            "Rank": int(sh2.iat[rr, 35]),
        })
    final_df = pd.DataFrame(final_rows).sort_values("Rank")

    return {
        "sheet_names": xls.sheet_names,
        "linguistic_scale_df": scale_df,
        "code_to_it2": code_to_it2,
        "experts": experts,
        "criteria_weights_df": crit_weight_df,
        "normalized_matrices": norm_mats,
        "final_results_df": final_df
    }


def module_cocoso():
    st.header("📊 IT2TrFS–CoCoSo (Excel-driven)")
    st.write("Upload your CoCoSo Excel file (same structure as your sample). The app will read the inputs and show the CoCoSo outputs (weights, normalized matrices, SBi/PBi, K, Rank).")

    # Use uploader; also allow using the mounted sample path if running locally in same environment.
    uploaded = st.file_uploader("Upload IT2TrFS-CoCoSo Excel (.xlsx)", type=["xlsx"])

    # If no upload, try to use the sample path you provided in this chat environment
    sample_path = "/mnt/data/d57c9134-88cf-4ba8-916f-735b7f628342.xlsx"

    data = None
    try:
        if uploaded is not None:
            data = parse_cocoso_excel(uploaded)
        else:
            # fallback: sample
            data = parse_cocoso_excel(sample_path)
            st.info("Using the provided sample Excel (mounted in the environment). Upload your own file to replace it.")
    except Exception as e:
        st.error(f"Failed to parse Excel: {e}")
        return

    # quick overview
    n_exp = len(data["experts"])
    n_crit = len(data["criteria_weights_df"])
    n_alt = len(data["final_results_df"])
    c1, c2, c3 = st.columns(3)
    c1.metric("Experts detected", n_exp)
    c2.metric("Criteria detected", n_crit)
    c3.metric("Alternatives detected", n_alt)

    with st.expander("Linguistic scale read from Excel", expanded=False):
        df = data["linguistic_scale_df"].copy()
        df["IT2TrFS (formatted)"] = df["IT2TrFS"].apply(lambda x: format_it2(x))
        st.dataframe(df[["Linguistic Attribute", "Code", "IT2TrFS (formatted)"]], use_container_width=True, hide_index=True)

    with st.expander("Expert decision matrices (linguistic codes) read from Excel", expanded=False):
        for ex in data["experts"]:
            st.subheader(ex["expert_label"])
            meta = pd.DataFrame({"Criterion": ex["criteria"], "Optimization": ex["optimization"]})
            st.dataframe(meta, use_container_width=True, hide_index=True)
            st.dataframe(ex["linguistic_matrix"], use_container_width=True)

    st.subheader("Criteria weights (from Excel)")
    st.dataframe(
        data["criteria_weights_df"].style.format({"Weight": "{:.6f}"}),
        use_container_width=True,
        hide_index=True
    )

    with st.expander("Normalized IT2TrFS decision matrices (from Excel)", expanded=False):
        # show one by one
        for alt, df_alt in data["normalized_matrices"].items():
            st.markdown(f"**{alt}**")
            show = df_alt.copy()
            show["IT2TrFS"] = show["IT2TrFS"].apply(lambda x: "N/A" if x is None else format_it2(x))
            st.dataframe(show, use_container_width=True)

    st.subheader("Final CoCoSo results (from Excel)")
    st.dataframe(
        data["final_results_df"].style.format({
            "Crisp_SBi": "{:.6f}",
            "Crisp_PBi": "{:.6f}",
            "Kia": "{:.6f}",
            "Kib": "{:.6f}",
            "Kic": "{:.6f}",
            "K": "{:.6f}",
        }),
        use_container_width=True,
        hide_index=True
    )

    # optional chart
    with st.expander("Ranking chart", expanded=False):
        df = data["final_results_df"].copy()
        fig, ax = plt.subplots(figsize=(9, 5))
        ax.bar(df["Technology"], df["K"])
        ax.set_ylabel("K (CoCoSo final score)")
        ax.set_title("IT2TrFS-CoCoSo Final Scores")
        plt.xticks(rotation=45, ha="right")
        st.pyplot(fig)


# ============================================================
# WINGS module UI (your original app flow moved into a function)
# ============================================================

def module_wings():
    st.header("📊 IT2TrFS–WINGS")
    st.write("""
    This tool implements the Interval Type-2 Trapezoidal Fuzzy Sets Weighted Influence Non-linear Gauge System (IT2TrFS WINGS)
    method for analyzing systems with interrelated components under uncertainty, incorporating input from multiple experts.
    """)

    tab_howto, tab_analysis = st.tabs(["📘 How to Use", "📊 Analysis"])

    with tab_howto:
        st.markdown("""
        ### Overview
        The IT2TrFS WINGS method helps analyze complex systems with interrelated components,
        handling uncertainty using Interval Type-2 Trapezoidal Fuzzy Sets (IT2TrFSs).
        """)

        with st.expander("Linguistic Terms Reference"):
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Strength/Relevance Terms**")
                strength_df = pd.DataFrame([
                    {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                    for abbr, it2 in LINGUISTIC_TERMS["strength"].items()
                ])
                st.dataframe(strength_df, hide_index=True, use_container_width=True)
            with col2:
                st.write("**Influence Terms**")
                infl_df = pd.DataFrame([
                    {"Abbreviation": abbr, "Full Form": FULL_FORMS[abbr], "IT2TrFS Interval": format_it2(it2)}
                    for abbr, it2 in LINGUISTIC_TERMS["influence"].items()
                ])
                st.dataframe(infl_df, hide_index=True, use_container_width=True)

    with tab_analysis:
        with st.sidebar:
            st.subheader("⚙️ WINGS Configuration")
            n_components = st.number_input("Number of Components", min_value=2, max_value=25, value=3)
            n_experts = st.number_input("Number of Experts", min_value=1, max_value=15, value=1)

            component_names = []
            for i in range(n_components):
                component_names.append(st.text_input(f"Name of Component {i+1}", value=f"C{i+1}", key=f"w_comp_{i}"))

            expert_weights = None
            if n_experts > 1:
                st.markdown("---")
                st.write("Assign weights to each expert (must sum to 1.0):")
                weights = []
                total = 0.0
                for i in range(n_experts):
                    max_val = min(1.0, 1.0 - total + (1.0/n_experts))
                    w = st.number_input(
                        f"Weight for Expert {i+1}",
                        min_value=0.0, max_value=max_val,
                        value=1.0/n_experts, step=0.01, format="%.2f",
                        key=f"w_expw_{i}"
                    )
                    weights.append(w)
                    total += w
                st.write(f"**Current total:** {total:.2f}/1.0")
                if abs(total - 1.0) > 0.001:
                    st.error("Weights must sum to 1.0.")
                    st.stop()
                expert_weights = weights

            st.markdown("---")
            st.info("Use abbreviations for strength and influence assessments.")

        if 'wings_experts_data' not in st.session_state:
            st.session_state.wings_experts_data = {}

        for expert_idx in range(n_experts):
            if expert_idx not in st.session_state.wings_experts_data:
                st.session_state.wings_experts_data[expert_idx] = {
                    'strengths_linguistic': ["HR" for _ in range(n_components)],
                    'influence_matrix_linguistic': [["ELI" for _ in range(n_components)] for _ in range(n_components)]
                }
            else:
                if len(st.session_state.wings_experts_data[expert_idx]['strengths_linguistic']) != n_components:
                    st.session_state.wings_experts_data[expert_idx]['strengths_linguistic'] = ["HR" for _ in range(n_components)]
                if len(st.session_state.wings_experts_data[expert_idx]['influence_matrix_linguistic']) != n_components:
                    st.session_state.wings_experts_data[expert_idx]['influence_matrix_linguistic'] = [["ELI" for _ in range(n_components)] for _ in range(n_components)]

        st.subheader("👨‍💼 Expert Input" if n_experts > 1 else "👨‍💼 Data Input")

        expert_tabs = st.tabs([f"Expert {i+1}" for i in range(n_experts)]) if n_experts > 1 else [st.container()]

        strengths_list = []
        influence_matrices_list = []

        for expert_idx in range(n_experts):
            tab = expert_tabs[expert_idx] if n_experts > 1 else expert_tabs[0]
            with tab:
                if n_experts > 1:
                    st.markdown(f"**Expert {expert_idx+1}**")
                    if expert_weights:
                        st.caption(f"Weight: {expert_weights[expert_idx]:.2f}")

                st.write("**Component Strengths/Relevance**")
                strengths = []
                cols = st.columns(n_components)
                for i in range(n_components):
                    with cols[i]:
                        cur = st.session_state.wings_experts_data[expert_idx]['strengths_linguistic'][i]
                        strength_term = st.selectbox(
                            component_names[i],
                            options=list(LINGUISTIC_TERMS["strength"].keys()),
                            index=list(LINGUISTIC_TERMS["strength"].keys()).index(cur),
                            key=f"w_strength_{expert_idx}_{i}"
                        )
                        st.session_state.wings_experts_data[expert_idx]['strengths_linguistic'][i] = strength_term
                        strengths.append(LINGUISTIC_TERMS["strength"][strength_term])

                st.write("**Influence Matrix** (row influences column)")
                influence_matrix = [[None]*n_components for _ in range(n_components)]
                for i in range(n_components):
                    row_cols = st.columns(n_components)
                    for j in range(n_components):
                        with row_cols[j]:
                            if i == j:
                                st.markdown("—")
                                influence_matrix[i][j] = zero_it2()  # placeholder not used for diag
                            else:
                                cur = st.session_state.wings_experts_data[expert_idx]['influence_matrix_linguistic'][i][j]
                                inf = st.selectbox(
                                    f"{component_names[i]}→{component_names[j]}",
                                    options=list(LINGUISTIC_TERMS["influence"].keys()),
                                    index=list(LINGUISTIC_TERMS["influence"].keys()).index(cur),
                                    key=f"w_inf_{expert_idx}_{i}_{j}",
                                    label_visibility="collapsed"
                                )
                                st.session_state.wings_experts_data[expert_idx]['influence_matrix_linguistic'][i][j] = inf
                                influence_matrix[i][j] = LINGUISTIC_TERMS["influence"][inf]

            strengths_list.append(strengths)
            influence_matrices_list.append(influence_matrix)

        if st.button("🚀 Run IT2TrFS WINGS Analysis", type="primary", use_container_width=True):
            with st.spinner("Calculating..."):
                results = wings_method_experts(strengths_list, influence_matrices_list, expert_weights)

            st.success("Analysis Complete!")

            tab1, tab3, tab4, tab5, tab6, tab7 = st.tabs([
                "🔗 Flowchart", "🧮 IT2TrFS Matrices", "📊 Results",
                "🏷️ Component Classification", "📈 Visualization", "📤 Export"
            ])

            with tab1:
                if n_experts > 1:
                    for ei in range(n_experts):
                        st.subheader(f"Expert {ei+1}")
                        st.graphviz_chart(generate_flowchart_for_expert(
                            st.session_state.wings_experts_data[ei], component_names, ei
                        ), use_container_width=True)
                else:
                    st.graphviz_chart(generate_flowchart_for_expert(
                        st.session_state.wings_experts_data[0], component_names
                    ), use_container_width=True)

            with tab3:
                st.subheader("Average SIDRM")
                st.dataframe(format_it2_df(results['average_sidrm'], component_names, component_names), use_container_width=True)
                st.subheader("Normalized Matrix Z")
                st.dataframe(format_it2_df(results['normalized_matrix'], component_names, component_names), use_container_width=True)
                st.subheader("Total Matrix T")
                st.dataframe(format_it2_df(results['total_matrix'], component_names, component_names), use_container_width=True)

            with tab4:
                res_df = pd.DataFrame({
                    "Component": component_names,
                    "TI (defuzz)": results["total_impact_defuzz"],
                    "TR (defuzz)": results["total_receptivity_defuzz"],
                    "Engagement (defuzz)": results["engagement_defuzz"],
                    "Role (defuzz)": results["role_defuzz"],
                }).sort_values("Engagement (defuzz)", ascending=False).reset_index(drop=True)
                res_df["Rank (by engagement)"] = np.arange(1, len(res_df)+1)
                st.dataframe(res_df.style.format(precision=6), use_container_width=True, hide_index=True)

            with tab5:
                class_df = pd.DataFrame({
                    "Component": component_names,
                    "Type": ["Cause" if r > 0 else "Effect" for r in results["role_defuzz"]],
                    "Role (defuzz)": results["role_defuzz"],
                    "Engagement (defuzz)": results["engagement_defuzz"],
                })
                st.dataframe(class_df.style.format({"Role (defuzz)": "{:.6f}", "Engagement (defuzz)": "{:.6f}"}),
                             use_container_width=True, hide_index=True)

            with tab6:
                fig, ax = plt.subplots(figsize=(9, 6))
                for i, name in enumerate(component_names):
                    ax.scatter(results["engagement_defuzz"][i], results["role_defuzz"][i], s=120)
                    ax.annotate(name, (results["engagement_defuzz"][i], results["role_defuzz"][i]), xytext=(6, 6), textcoords="offset points")
                ax.axhline(0, linestyle="--")
                ax.set_xlabel("Engagement (defuzz)")
                ax.set_ylabel("Role (defuzz)")
                ax.set_title("Engagement vs Role")
                st.pyplot(fig)

            with tab7:
                doc = create_word_report(results, component_names, n_experts, expert_weights)
                st.markdown(get_word_download_link(doc), unsafe_allow_html=True)


# ============================================================
# MAIN APP: sidebar navigation with two modules
# ============================================================

def main():
    st.set_page_config(page_title="IT2TrFS Toolkit (WINGS + CoCoSo)", layout="wide", page_icon="📊")

    st.title("🧰 IT2TrFS Toolkit")
    st.sidebar.header("Navigation")
    page = st.sidebar.radio("Choose a Model", ["IT2TrFS–WINGS", "IT2TrFS–CoCoSo"], index=0)
    st.sidebar.markdown("---")
    st.sidebar.info("Both modules are independent. Use the left menu to switch.")

    if page == "IT2TrFS–WINGS":
        module_wings()
    else:
        module_cocoso()

if __name__ == "__main__":
    main()
